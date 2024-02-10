# developing in progress .......

import openai
import gradio as gr
import PyPDF2
import docx
from docx import Document
from gradio import CSVLogger
import os
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph
from reportlab.lib.units import inch
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
    


openai.api_key = ""

messages = []

latest_resume_text = ""

default_text_1 = "Hello, I will share a candidate's resume with you. Please provide a summary for me" 
                
default_text_2 = "Hello, I will share a candidate's resume with you. Please desensitize the resume by removing or redacting any personal contact information like the phone number, home address, and email address. Also remove or mask any private details like ID numbers, exact dates for education and work history, and compensation details. Double check that no other private information is included. The goal is to remove personally identifying details while still allowing this candidate's professional qualifications to be visible to potential employers."

# Chinese version
default_text_3 = "请帮我阅读候选人简历并总结出以下几个部分。第一个板块是“个人信息”，请总结出此后选人的姓名，性别，工作经验（多少年），最高学历，毕业院校，专业，和毕业时间。第二个部分是此候选人的本人评价。第三个部分是具体得工作经历。第四个部分是做过的项目经验。请根据你对于这位候选人简历的最细致的阅读排出以上几个部分。每个部分由自己的标题：五个标题为 个人信息，本人评价，工作经历，和项目经验。"


def set_text_1():
    return default_text_1
def set_text_2():
    return default_text_2
def set_text_3():
    return default_text_3

# extracting text from pdf
def extract_text_from_pdf(file):
    reader = PyPDF2.PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

# extracting text from doc
def extract_text_from_docx(file):

    doc = docx.Document(file)
    all_text = []

    for para in doc.paragraphs:
        # Extract the text from the current paragraph
        paragraph_text = para.text
        all_text.append(paragraph_text)

    combined_text = " ".join(all_text)
    return combined_text

def handle_file_upload(uploaded_files):
    combined_text = ""
    if uploaded_files is None:
        return ""
    for uploaded_file in uploaded_files:
        file_type = uploaded_file.name.split('.')[-1].lower()
        if file_type == 'pdf':
            combined_text += extract_text_from_pdf(uploaded_file)
        elif file_type in ['docx', 'doc']:
            combined_text += extract_text_from_docx(uploaded_file)
        else:
            combined_text += "Unsupported file format, please upload a PDF or Word document.\n"
    return combined_text

    
def clear_inputs():
    return "", None

def log_conversation(input, reply):
    with open("log.csv", "a") as log_file:
        log_file.write(f"user_input: {input},\n\n Chatgpt: {reply}\n\n")


import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def is_chinese_char(char):
    """Check if a given character is a Chinese character."""
    return '\u4e00' <= char <= '\u9fff'

def count_chinese_chars(text):
    """Count the number of Chinese characters in a string."""
    return sum(is_chinese_char(char) for char in text)


def generate_resume_document(latest_resume_text):
    # Define keywords that indicate the start of a new section
    section_keywords = ["个人信息", "本人评价", "工作经历", "项目经验"]

    # Create a new Document
    doc = Document()

    # Add a title to the document
    doc.add_heading('候选人简历', level=0)

    # Check if the text is empty
    if not latest_resume_text.strip():
        doc.add_paragraph("No resume text found")
        return None

    # Remove non-word characters from the text except for spaces and new lines
    # Remove non-word characters from the text except for "-", ",", and "."
    cleaned_text = re.sub(r'[^\w\s\-，。]', '', latest_resume_text)


    # Split the resume text into lines
    lines = cleaned_text.split('\n')

    # Variables to store the current section title and its content
    current_section = None
    current_content = []

    # Function to add a section to the document with bullet points
    def add_section_to_doc(section, content):
        if section and content:
            doc.add_heading(section, level=2)
            for line in content:
                # Check if line has less than 20 characters (excluding spaces)
                chinese_char_count = count_chinese_chars(line)
                if section == "个人信息":
                # Combine all content lines into a single line separated by a semicolon
                    combined_content = '；'.join(content)  # This uses a semicolon as a separator
                    doc.add_paragraph(combined_content)
                    break
                if chinese_char_count < 5:
                    # If the line has less than 20 Chinese characters, make it a sub-sub-header
                    doc.add_heading(line, level=3)
                else:
                    # Otherwise, add it as a bullet point
                    p = doc.add_paragraph()
                    p.add_run(line).bold = True  # Optional: make bullet points bold
            

    for line in lines:
        # Check for section titles (considering case-insensitivity)
        if any(line.strip().upper() == keyword for keyword in section_keywords):
            # Add the previous section to the document before starting a new one
            add_section_to_doc(current_section, current_content)
            current_section = line.strip().title()  # Title Case for headings
            current_content = []
        else:
            # Clean line and add to current content
            clean_line = line.strip()
            if clean_line:  # Ignore empty lines
                # Append clean_line to current_content, removing any additional internal line breaks
                current_content.append(re.sub(r'\s+', ' ', clean_line))

    # Add the last section to the document
    add_section_to_doc(current_section, current_content)


    # Save the document
    word_filename = "generated_resume.docx"
    doc.save(word_filename)
    return word_filename





def CustomChatGPT(user_input, uploaded_file):
    global messages, latest_resume_text  # Declare latest_resume_text as global if it's used globally

    resume_text = handle_file_upload(uploaded_file)
    print("Resume text from the file:", resume_text)

    if resume_text:
        messages.append({"role": "system", "content": resume_text})

    messages.append({"role": "user", "content": user_input})
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=messages
    )
    ChatGPT_reply = response.choices[0].message.content
    messages.append({"role": "assistant", "content": ChatGPT_reply})

    log_conversation(user_input, ChatGPT_reply)

    latest_resume_text = ChatGPT_reply

    if "请帮我阅读候选人简历并总结出以下几个部分" in user_input:
        word_filename = generate_resume_document(latest_resume_text)
        ChatGPT_reply = "Resume generated successfully. Please download the file using the link below."
    else:
        doc = Document()
        doc.add_paragraph(ChatGPT_reply)
        word_filename = "generated_response.docx"
        doc.save(word_filename)

    return ChatGPT_reply, word_filename



def get_log_file():
    return "log.csv"
def read_log_file():
    if os.path.exists("log.csv"):
        with open("log.csv", "r") as file:
            return file.read()
    else:
        return "Log file is empty or doesn't exist."
def clear_log_file():
    if not os.path.exists("log.csv"):
        return "file not exist"
    else:
        with open("log.csv", "w") as log_file:
            log_file.write("") 
        return "Log cleared"

def main():
    custom_css = """
    <style>
    .custom-button .mdc-button {
        background-color: #90EE90; /* Light Green */
        color: black;
        border: none;
        border-radius: 6px;
        padding: 10px 20px;
        margin: 8px 0;
        cursor: pointer;
        font-size: 14px;
        text-align: center;
    }
    .custom-button .mdc-button:hover {
        background-color: #76b476; /* Darker shade for hover effect */
    }
    </style>
    """

    with gr.Blocks(css=custom_css) as demo:
        gr.Markdown("# Resume Analysis Tool", elem_classes=["centered-title"])

        with gr.Row():
            with gr.Column():
                text_input = gr.Textbox()
                file_input = gr.File(file_count="multiple", label="Upload Resume")

                with gr.Row():
                    btn1 = gr.Button("Summary", elem_classes=["custom-button"])
                    btn2 = gr.Button("Desensitization", elem_classes=["custom-button"])
                    btn3 = gr.Button("简历生成", elem_classes=["custom-button"])


                    btn1.click(fn=set_text_1, inputs=[], outputs=text_input)
                    btn2.click(fn=set_text_2, inputs=[], outputs=text_input)
                    btn3.click(fn=set_text_3, inputs=[], outputs=text_input)
                    

                with gr.Row():
                    submit_btn = gr.Button("Submit", elem_classes=["custom-button"])
                    clear_btn = gr.Button("Clear", elem_classes=["custom-button"])
                    clear_btn.click(fn=clear_inputs, inputs=[], outputs=[text_input, file_input])

            with gr.Column():
                with gr.Row():
                    output_text = gr.Textbox(label="ChatGPT Reply", interactive=True, lines=1)
                with gr.Row():
                    output_word = gr.File(label="Download Word File")
                    # output_pdf = gr.File(label="Download PDF File")

                    submit_btn.click(fn=CustomChatGPT, inputs=[text_input, file_input], outputs=[output_text, output_word])

                with gr.Row():
                    log_text = gr.Textbox(label="Log Content", interactive=True, lines=1)

                with gr.Row():
                    view_log_button = gr.Button("View Log", elem_classes=["custom-button"])
                    download_log_button = gr.Button("Download Log File", elem_classes=["custom-button"])
                    clear_log_button = gr.Button("Clear Log File", elem_classes=["custom-button"])

                    view_log_button.click(fn=read_log_file, inputs=[], outputs=log_text)
                    download_log_button.click(fn=get_log_file, inputs=[], outputs=[])
                    clear_log_button.click(fn=clear_log_file, inputs=[], outputs=[])

                
    demo.launch(share=True)


if __name__ == "__main__":
    main()
