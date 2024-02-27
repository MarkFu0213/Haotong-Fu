import openai
import gradio as gr
import PyPDF2
import docx
from docx import Document
import os
from docx import Document
from docx.shared import Pt

from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_ALIGN_PARAGRAPH
  


openai.api_key = ""

messages = []


latest_resume_text = ""
desired_position = ""

# summarizing resume content
default_text_1 = "请帮我总结我上传的简历" 
# desensitization
default_text_2 = "请帮我把我上传的简历进行脱敏处理。"
# output specific resume format
default_text_3 = "请帮我阅读候选人简历并总结出以下几个部分。第一个板块是“个人信息”，请总结出此后选人的姓名，性别，工作经验（多少年），最高学历，毕业院校，专业，和毕业时间。第二个部分是此候选人的本人评价，放进一个自然段里。第三个部分是具体得工作经历。第四个部分是做过的项目经验。请根据你对于这位候选人简历的最细致的阅读排出以上几个部分。每个部分由自己的标题：五个标题为 个人信息，本人评价，工作经历，和项目经验。请把每一个小项写得细致一点，并且用数字排序！"

empty_text = ""

def set_empty_text():
    return empty_text

def set_text_1():
    return default_text_1
def set_text_2():
    return default_text_2
def set_text_3():
    return default_text_3

# # extracting text from pdf
# def extract_text_from_pdf(file):
#     reader = PyPDF2.PdfReader(file)
#     text = ""
#     for page in reader.pages:
#         text += page.extract_text()
#     return text



#write a function that extract text from a pdf that could contain multiple pages
def extract_text_from_pdf(file):
    reader = PyPDF2.PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

# extracting text from doc
def extract_text_from_docx(file):

    doc = docx.Document(file)
    all_text = []

    for para in doc.paragraphs:
        # Extract the text from the current paragraph
        paragraph_text = para.text
        all_text.append(paragraph_text)

    combined_text = " ".join(all_text)
    return combined_text

def handle_file_upload(uploaded_files):
    combined_text = ""
    if uploaded_files is None:
        return ""
    for uploaded_file in uploaded_files:
        file_type = uploaded_file.name.split('.')[-1].lower()
        if file_type == 'pdf':
            combined_text += extract_text_from_pdf(uploaded_file)
        elif file_type in ['docx', 'doc']:
            combined_text += extract_text_from_docx(uploaded_file)
        else:
            combined_text += "Unsupported file format, please upload a PDF or Word document.\n"
    return combined_text

    
def clear_inputs():
    messages = []
    return "", None



import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def is_chinese_char(char):
    """Check if a given character is a Chinese character."""
    return '\u4e00' <= char <= '\u9fff'

def count_chinese_chars(text):
    """Count the number of Chinese characters in a string."""
    return sum(is_chinese_char(char) for char in text)


def generate_resume_document(latest_resume_text):
    # Define keywords that indicate the start of a new section
    section_keywords = ["个人信息", "本人评价", "工作经历", "项目经验"]

    # Create a new Document
    doc = Document()

    # Add a title to the document
    heading = doc.add_heading('候选人简历', level=0)
    run = heading.add_run()
    run.underline = WD_UNDERLINE.SINGLE

# Center-align the heading
    paragraph_format = heading.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Check if the text is empty
    if not latest_resume_text.strip():
        doc.add_paragraph("No resume text found")
        return None

    # Remove non-word characters from the text except for spaces and new lines
    # Remove non-word characters from the text except for "-", ",", "." and numbers

    cleaned_text = re.sub(r'[^\w\s\-\,\.\d]', '', latest_resume_text)


    # Split the resume text into lines
    lines = cleaned_text.split('\n')

    # Variables to store the current section title and its content
    current_section = None
    current_content = []

    # Function to add a section to the document with bullet points
    def add_section_to_doc(section, content):
        global name
        if section and content:
            sec = doc.add_heading(section, level=0)
            for run in sec.runs:
                run.font.size = Pt(14)
   
        
            
            if section == "个人信息":
                # Combine all content lines into a single line separated by a semicolon
                combined_content = '；'.join(content)
                p = doc.add_paragraph()  # Create a new paragraph for combined content

                last_index = 0  # Keep track of the last index processed
                for keyword in ["姓名", "性别", "工作经验", "最高学历", "毕业院校", "专业", "毕业时间"]:
                    if keyword in combined_content:
                        start_index = combined_content.index(keyword, last_index)

                        # Add text before the keyword as a normal run
                        p.add_run(combined_content[last_index:start_index])

                        # Add the keyword and the colon as a bold run
                        bold_run = p.add_run(keyword + '：')
                        bold_run.bold = True

                        # Update the last index processed
                        last_index = start_index + len(keyword)

                    if keyword == "姓名":
                        name_start_index = last_index
                        try:
                            # Try to find the start index of the next keyword "性别"
                            name_end_index = combined_content.index("性别", name_start_index)
                            # Store the content into name variable
                            name = combined_content[name_start_index:name_end_index]
                            # Remove any non-Chinese characters
                            name = re.sub(r'[^\u4e00-\u9fff]', '', name)
                        except ValueError:
                            # Handle the case where "性别" is not found
                            name = "需手动填写"

                        
                        


                        

                # Add any remaining text after the last keyword
                p.add_run(combined_content[last_index:])
                


                
            else:
                for line in content:
                    # # Count Chinese characters in the line
                    # chinese_char_count = count_chinese_chars(line)
                    # if chinese_char_count < 5:
                    #     doc.add_heading(line, level=0)
                    #     p = doc.add_paragraph()
                    #     p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    #     for run in sec.runs:
                    #         run.font.size = Pt(14)

                    # else:
                   
                    doc.add_paragraph(line)
                    
            

    for line in lines:
        # Check for section titles (considering case-insensitivity)
        if any(line.strip().upper() == keyword for keyword in section_keywords):
            # Add the previous section to the document before starting a new one
            add_section_to_doc(current_section, current_content)
            
            
            

            current_section = line.strip().title()  # Title Case for headings
            current_content = []
        else:
            # Clean line and add to current content
            clean_line = line.strip()
            if clean_line:  # Ignore empty lines
                # Append clean_line to current_content, removing any additional internal line breaks
                current_content.append(re.sub(r'\s+', ' ', clean_line))

    # Add the last section to the document
    add_section_to_doc(current_section, current_content)
    
    # Save the document
    word_filename = "博网科技-" + name + "-" + desired_position + ".docx"
    doc.save(word_filename)
    return word_filename




def log_conversation(input, reply):
    with open("log.csv", "a") as log_file:
        log_file.write(f"user_input: {input},\n\n Chatgpt: {reply}\n\n")



def CustomChatGPT(user_input, uploaded_file):
    global messages, latest_resume_text, desired_position, name  # Declare latest_resume_text as global if it's used globally
    resume_text = ""
    resume_text = handle_file_upload(uploaded_file)
    print("Resume text from the file:", resume_text)

    if resume_text:
        messages.append({"role": "system", "content": resume_text})
        # if the resume text contains "期望职位", extract the desired position and store it in desired_position
        if "期望职位" in resume_text:
            match = re.search(r'期望职位：(.+?)\s', resume_text)
            if match:
                desired_position = match.group(1)
                # Remove non-Chinese characters
                desired_position = re.sub(r'[^\u4e00-\u9fff]', '', desired_position)
            else:
                desired_position = "需手动填写"
                
            

    messages.append({"role": "user", "content": user_input})
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=messages
    )
    ChatGPT_reply = response.choices[0].message.content
    messages.append({"role": "assistant", "content": ChatGPT_reply})

    combined_input = user_input + " " + resume_text
    log_conversation(combined_input, ChatGPT_reply)

    latest_resume_text = ChatGPT_reply

    if "请帮我阅读候选人简历" in user_input:
        word_filename = generate_resume_document(latest_resume_text)
        ChatGPT_reply = "Resume generated successfully. Please download the file using the link below."
        messages = []
    else:
        doc = Document()
        doc.add_paragraph(ChatGPT_reply)
        word_filename = "回复.docx"
        doc.save(word_filename)




        # reset the messages to start a new conversation
        messages = []

    return ChatGPT_reply, word_filename



def get_log_file():
    log_content = read_log_file()
    if log_content == "Log file is empty or doesn't exist.":
        return None  # Or handle the empty log case as you prefer
    # strip all leading/trailing whitespaces and newlines
    log_content = log_content.strip()

    # Create a new Word document
    doc = Document()
    doc.add_heading('对话记录', 0)

    # Add log content to the Word document
    for line in log_content.split('\n'):
        doc.add_paragraph(line)

    # Save the document
    doc_path = "chat_log.docx"
    doc.save(doc_path)

    # Return the path to the saved document for Gradio to handle the download
    return doc_path







def read_log_file():
    if os.path.exists("log.csv"):
        with open("log.csv", "r") as file:
            return file.read()
    else:
        return "Log file is empty or doesn't exist."
def clear_log_file():
    if not os.path.exists("log.csv"):
        return "file not exist"
    else:
        with open("log.csv", "w") as log_file:
            log_file.write("") 
        return "Log cleared"
    

def handle_btn1_click(file_input):
    return CustomChatGPT(default_text_1, file_input)

def handle_btn2_click(file_input):
    return CustomChatGPT(default_text_2, file_input)

def handle_btn3_click(file_input):
    return CustomChatGPT(default_text_3, file_input)

def main():
    custom_css = """
    <style>
    .custom-button .mdc-button {
        background-color: #90EE90; /* Light Green */
        color: black;
        border: none;
        border-radius: 6px;
        padding: 10px 20px;
        margin: 8px 0;
        cursor: pointer;
        font-size: 14px;
        text-align: center;
    }
    .custom-button .mdc-button:hover {
        background-color: #76b476; /* Darker shade for hover effect */
    }
    </style>
    """

    with gr.Blocks(css=custom_css) as demo:
        gr.Markdown("# 简历分析系统", elem_classes=["centered-title"])

        with gr.Row():
            with gr.Column():
                # Inputs
                text_input = gr.Textbox()
                file_input = gr.File(file_count="multiple", label="Upload Resume")
                with gr.Row():
                # Buttons for resume operations
                    btn1 = gr.Button("总结简历", elem_classes=["custom-button"])
                    btn2 = gr.Button("脱敏处理", elem_classes=["custom-button"])
                    btn3 = gr.Button("简历生成", elem_classes=["custom-button"])
                with gr.Row():

                # Other buttons
                    submit_btn = gr.Button("提交", elem_classes=["custom-button"])
                    clear_btn = gr.Button("清除", elem_classes=["custom-button"])
                    clear_btn.click(fn=set_empty_text, inputs=[], outputs=[text_input])

            with gr.Column():
                # Outputs
                output_text = gr.Textbox(label="ChatGPT回复", interactive=True, lines=1)
                output_word = gr.File(label="下载word文件")

                # Click events for resume operation buttons
                btn1.click(fn=handle_btn1_click, inputs=[file_input], outputs=[output_text, output_word])
                btn2.click(fn=handle_btn2_click, inputs=[file_input], outputs=[output_text, output_word])
                btn3.click(fn=handle_btn3_click, inputs=[file_input], outputs=[output_text, output_word])

                # Click event for submitting the chat
                submit_btn.click(fn=CustomChatGPT, inputs=[text_input, file_input], outputs=[output_text, output_word])

                # Log-related components and buttons
                log_text = gr.Textbox(label="Log Content", interactive=True, lines=1)
                output_docx = gr.File(label="下载log文件")  # For downloading the log file
      

                
                with gr.Row():
                    view_log_button = gr.Button("对话记录", elem_classes=["custom-button"])
                    download_log_button = gr.Button("下载对话记录", elem_classes=["custom-button"])
                    clear_log_button = gr.Button("清除对话记录", elem_classes=["custom-button"])
                    
                    view_log_button.click(fn=read_log_file, inputs=[], outputs=log_text)
                    
                    download_log_button.click(fn=get_log_file, inputs=[], outputs=[output_docx])
                    clear_log_button.click(fn=clear_log_file, inputs=[], outputs=[])



    demo.launch()



if __name__ == "__main__":
    main()
